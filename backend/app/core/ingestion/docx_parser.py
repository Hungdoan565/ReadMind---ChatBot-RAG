"""
DOCX parser using python-docx.
Extracts paragraphs and tables.
"""

import logging
import uuid
from pathlib import Path
from typing import List

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


def parse_docx(file_path: str | Path, doc_id: str | None = None) -> List[Document]:
    """Parse a .docx file into Documents."""
    import docx

    file_path = Path(file_path)
    if not doc_id:
        doc_id = str(uuid.uuid4())

    doc = docx.Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    return [
        Document(
            page_content=full_text,
            metadata={
                "source": file_path.name,
                "doc_id": doc_id,
                "file_type": "docx",
            },
        )
    ]


def parse_docx_bytes(
    content: bytes, filename: str, doc_id: str | None = None, room_code: str | None = None,
) -> List[Document]:
    """Parse DOCX from bytes."""
    import io
    import docx
    
    if not doc_id:
        doc_id = str(uuid.uuid4())
    
    docx_doc = docx.Document(io.BytesIO(content))
    paragraphs = [p.text for p in docx_doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    
    return [
        Document(
            page_content=full_text,
            metadata={
                "source": filename,
                "doc_id": doc_id,
                "file_type": "docx",
                "room_code": room_code or "",
            },
        )
    ]
